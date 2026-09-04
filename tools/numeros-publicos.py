#!/usr/bin/env python3
"""
Os números que os documentos de fora dizem sobre os projectos batem certo?

PORQUE E QUE ISTO EXISTE

A 4 de setembro de 2026 o CV que ia para candidaturas dizia: cadence 44 testes,
framebudget 7.3 KB e 24 testes, glaze 16.1 KB e 83 testes. Os números reais eram
96, 11.7 KB, 58, 17.2 KB e 91. Cinco números errados num documento cujo objectivo
é ser lido por quem contrata, e cada um deles era uma pergunta de entrevista à
espera de acontecer.

Nenhum foi inventado quando foi escrito — eram todos verdade no dia. Depois os
projectos cresceram e o CV não. É o mesmo defeito que este repositório passou
três semanas a apagar de dentro dos READMEs, com a diferença de que dentro do
repositório há guardas em CI e fora não havia nada.

Isto é esse guarda, para fora. Lê os números que os documentos públicos afirmam e
compara-os com o que os próprios repositórios já garantem sobre si.

O QUE ESTE PROGRAMA NÃO FAZ

Não conta testes nem mede bundles. Isso já é feito, e melhor, pelos `contagem` e
`tamanho` de cada repositório, que correm em CI e falham a build. Se ele
recontasse por sua conta, passava a haver duas fontes para o mesmo número e um
dia discordavam — que é precisamente a doença. Lê os READMEs, que são o que os
guardas mantêm honesto.

    python3 tools/numeros-publicos.py
    python3 tools/numeros-publicos.py --cv "../../Rodrigo_Figueiredo_CV.docx"
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

RAIZ = Path(__file__).resolve().parent.parent
IRMAOS = RAIZ.parent


# ---------------------------------------------------------------- a verdade

def ler(caminho: Path) -> str | None:
    try:
        return caminho.read_text(encoding="utf-8")
    except OSError:
        return None


def numero(texto: str | None, padrao: str) -> str | None:
    if not texto:
        return None
    m = re.search(padrao, texto)
    return m.group(1) if m else None


def verdade() -> dict[str, dict[str, str]]:
    """
    O que cada repositório garante sobre si próprio, tirado do seu README —
    o mesmo sítio que o guarda desse repositório verifica em cada push.
    """
    out: dict[str, dict[str, str]] = {}

    fb = ler(IRMAOS / "framebudget" / "README.md")
    if fb:
        out["framebudget"] = {
            k: v for k, v in {
                "testes": numero(fb, r"(\d+) tests"),
                "kb": numero(fb, r"([\d.]+) KB minified"),
            }.items() if v
        }

    gl = ler(IRMAOS / "glaze" / "README.md")
    if gl:
        out["glaze"] = {
            k: v for k, v in {
                "testes": numero(gl, r"(\d+) tests, no browser, no GPU"),
                "kb": numero(gl, r"([\d.]+) KB minified"),
            }.items() if v
        }

    ca = ler(IRMAOS / "cadence" / "README.md")
    if ca:
        out["cadence"] = {
            k: v for k, v in {
                "testes": numero(ca, r"tests-(\d+)-"),
                "testes_backend": numero(ca, r"(\d+) backend tests"),
            }.items() if v
        }

    for nome in ("Tamagotchi", "Personal-Finance-Tracking-App"):
        t = ler(IRMAOS / nome / "README.md")
        if t:
            n = numero(t, r"(\d+) tests")
            if n:
                out[nome] = {"testes": n}

    return out


# ------------------------------------------------------- o que se diz por fora

def texto_do_docx(caminho: Path) -> str | None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        print("  ..  python-docx nao instalado; salto o CV "
              "(pip install python-docx)", file=sys.stderr)
        return None
    try:
        d = Document(str(caminho))
    except Exception as e:  # noqa: BLE001 - queremos o motivo, seja qual for
        print(f"  ..  nao consegui abrir {caminho.name}: {e}", file=sys.stderr)
        return None
    partes = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            partes += [c.text for c in row.cells]
    return "\n".join(partes)


# Cada alegação: (rótulo, regex com UM grupo, projecto, campo).
#
# Os padrões são propositadamente estreitos. Um padrão largo apanha o número
# errado e depois "passa" — que é pior do que não verificar, porque parece
# cobertura.
#
# Frases diferentes para os mesmos números, porque o CV e os posts do LinkedIn
# não escrevem a mesma coisa da mesma maneira.
ALEGACOES = [
    ("cadence: testes contra Postgres", r"(\d+) tests against real PostgreSQL",
     "cadence", "testes_backend"),
    ("cadence: testes (post)",
     r"(\d+) tests\. The backend ones run against real PostgreSQL",
     "cadence", "testes"),

    # Ancorado no nome do projecto: o CV não tem cabeçalhos, é uma secção só,
    # e sem a âncora o padrão do glaze apanhava o número do framebudget — que
    # foi o primeiro falso positivo que este programa deu.
    ("framebudget: tamanho (CV)",
     r"framebudget[\s\S]{0,400}?([\d.]+) KB minified", "framebudget", "kb"),
    ("framebudget: testes (CV)", r"(\d+) tests \(node:test\)",
     "framebudget", "testes"),
    ("framebudget: tamanho (post)",
     r"([\d.]+) KB, zero dependencies, \d+ tests", "framebudget", "kb"),
    ("framebudget: testes (post)",
     r"[\d.]+ KB, zero dependencies, (\d+) tests", "framebudget", "testes"),

    ("glaze: tamanho (CV)",
     r"glaze[\s\S]{0,400}?([\d.]+) KB minified", "glaze", "kb"),
    ("glaze: testes (CV)", r"(\d+) tests plus a browser page", "glaze", "testes"),
    ("glaze: tamanho (post)",
     r"([\d.]+) KB, zero dependencies, \d+ tests", "glaze", "kb"),
    ("glaze: testes (post)",
     r"[\d.]+ KB, zero dependencies, (\d+) tests", "glaze", "testes"),
    ("glaze: suite completa (post)", r"while all (\d+) tests", "glaze", "testes"),
]

# Um padrão sem grupo rebenta com IndexError a meio da verificação, que foi
# exactamente o que aconteceu da primeira vez. Verificar isto aqui custa nada e
# transforma um erro em tempo de execução num erro em tempo de arranque.
for _rot, _pad, _pr, _c in ALEGACOES:
    if re.compile(_pad).groups != 1:
        raise SystemExit(f"numeros-publicos: o padrão de {_rot!r} não tem exactamente um grupo")


def seccoes(texto: str) -> list[tuple[str, str]]:
    """
    Parte o documento por cabeçalhos `## `, para uma alegação sobre o glaze não
    ser verificada contra a frase do framebudget. Sem cabeçalhos, é tudo uma
    secção só — que é o caso do CV.
    """
    partes = re.split(r"^##\s+(.+)$", texto, flags=re.M)
    if len(partes) == 1:
        return [("", texto)]
    out = [("", partes[0])]
    for i in range(1, len(partes), 2):
        out.append((partes[i], partes[i + 1] if i + 1 < len(partes) else ""))
    return out


def verificar(nome_doc: str, texto: str, real: dict[str, dict[str, str]]) -> int:
    print(f"\n{nome_doc}")
    mal = 0
    vistos = 0
    for titulo, corpo in seccoes(texto):
        for rotulo, padrao, projecto, campo in ALEGACOES:
            # Se o documento tem secções, só verifico a alegação na secção do
            # projecto a que ela pertence.
            if titulo and projecto.lower() not in titulo.lower():
                continue
            m = re.search(padrao, corpo)
            if not m:
                continue
            vistos += 1
            diz = m.group(1)
            e = real.get(projecto, {}).get(campo)
            onde = f" [{titulo.strip()[:28]}]" if titulo else ""
            if e is None:
                print(f"  ..  {rotulo}{onde}: diz {diz}, e não sei o real "
                      f"(o README do {projecto} não está ao lado)")
                continue
            if diz == e:
                print(f"  ok  {rotulo}{onde}: {diz}")
            else:
                mal += 1
                print(f"  !!  {rotulo}{onde}: diz {diz}, é {e}")
    if not vistos:
        print("  ..  nenhuma das alegações conhecidas aparece neste documento")
    return mal


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--cv", default="../../Rodrigo_Figueiredo_CV.docx",
                    help="caminho do CV .docx (relativo a v1/)")
    ap.add_argument("--posts", default="docs/POSTS-LINKEDIN.md",
                    help="caminho do ficheiro de posts (relativo a v1/)")
    args = ap.parse_args()

    real = verdade()
    if not real:
        print("numeros-publicos: nao encontrei nenhum README dos irmãos.")
        print("  Corre isto na pasta a sério, com os repositórios ao lado.")
        return 0

    print("O que os repositórios garantem sobre si:")
    for proj, campos in sorted(real.items()):
        print("  " + proj + ": " + ", ".join(f"{k}={v}" for k, v in campos.items()))

    mal = 0
    posts = RAIZ / args.posts
    t = ler(posts)
    if t:
        mal += verificar(args.posts, t, real)

    cv = (RAIZ / args.cv).resolve()
    if cv.exists():
        t = texto_do_docx(cv)
        if t:
            mal += verificar(cv.name, t, real)
    else:
        print(f"\n  ..  {cv} nao existe — salto o CV")

    print()
    if mal:
        print(f"numeros-publicos: {mal} número(s) por corrigir lá fora.")
        return 1
    print("numeros-publicos: os documentos de fora dizem o mesmo que os repositórios.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
